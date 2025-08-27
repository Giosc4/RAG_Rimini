"""
Document Preprocessor con supporto Multimodale
Script per estrarre, pulire e consolidare documenti inclusi elementi multimodali
(immagini, tabelle, grafici) ottimizzato per il successivo chunking
"""

import os
import re
import json
import hashlib
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass, field, asdict
import logging
import unicodedata
import base64
from io import BytesIO

# Importazioni per gestire diversi tipi di file
try:
    from pypdf import PdfReader
    import fitz  # PyMuPDF per estrazione avanzata
except ImportError:
    print("pypdf/PyMuPDF non installato. Installa con: pip install pypdf PyMuPDF")
    PdfReader = None
    fitz = None

try:
    import chardet
except ImportError:
    print("chardet non installato. Installa con: pip install chardet")
    chardet = None

try:
    from docx import Document as DocxDocument
    from docx.table import Table
except ImportError:
    print("python-docx non installato. Installa con: pip install python-docx")
    DocxDocument = None

try:
    from bs4 import BeautifulSoup
except ImportError:
    print("beautifulsoup4 non installato. Installa con: pip install beautifulsoup4")
    BeautifulSoup = None

try:
    from PIL import Image
    import pytesseract  # Per OCR opzionale
except ImportError:
    print("PIL/pytesseract non installato. Installa con: pip install Pillow pytesseract")
    Image = None
    pytesseract = None

try:
    import pandas as pd
    import tabulate
except ImportError:
    print("pandas/tabulate non installato. Installa con: pip install pandas tabulate")
    pd = None
    tabulate = None

# Configurazione logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


@dataclass
class ExtractedElement:
    """Elemento estratto dal documento (testo, immagine, tabella, grafico)"""
    element_type: str  # 'text', 'image', 'table', 'chart'
    content: Any  # Contenuto dell'elemento
    position: Dict[str, Any] = field(default_factory=dict)  # Posizione nel documento
    metadata: Dict[str, Any] = field(default_factory=dict)
    element_id: str = ""
    page_number: Optional[int] = None
    
    def __post_init__(self):
        if not self.element_id:
            # Genera ID univoco per l'elemento
            content_hash = hashlib.md5(str(self.content).encode()).hexdigest()[:8]
            self.element_id = f"{self.element_type}_{content_hash}"


@dataclass
class ExtractedDocument:
    """Documento estratto e processato con elementi multimodali"""
    source_path: str
    file_type: str
    title: str
    content: str  # Testo principale consolidato
    elements: List[ExtractedElement] = field(default_factory=list)
    images: List[Dict[str, Any]] = field(default_factory=list)
    tables: List[Dict[str, Any]] = field(default_factory=list)
    sections: List[Dict[str, str]] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    extraction_date: str = field(default_factory=lambda: datetime.now().isoformat())
    word_count: int = 0
    char_count: int = 0
    line_count: int = 0
    image_count: int = 0
    table_count: int = 0


class MultimodalDocumentExtractor:
    """Estrae contenuto multimodale da diversi formati di file"""
    
    def __init__(self, extract_images: bool = True, extract_tables: bool = True,
                 ocr_images: bool = False, save_images: bool = True,
                 images_dir: str = "./extracted_images"):
        self.extract_images = extract_images
        self.extract_tables = extract_tables
        self.ocr_images = ocr_images
        self.save_images = save_images
        self.images_dir = Path(images_dir)
        
        if save_images:
            self.images_dir.mkdir(exist_ok=True)
        
        self.supported_formats = {
            '.pdf': self.extract_pdf_multimodal,
            '.txt': self.extract_text,
            '.md': self.extract_markdown,
            '.docx': self.extract_docx_multimodal,
            '.html': self.extract_html,
            '.json': self.extract_json
        }
    
    def extract(self, file_path: str) -> Optional[ExtractedDocument]:
        """Estrae contenuto multimodale dal file specificato"""
        path = Path(file_path)
        
        if not path.exists():
            logger.error(f"File non trovato: {file_path}")
            return None
        
        ext = path.suffix.lower()
        
        if ext not in self.supported_formats:
            logger.warning(f"Formato non supportato: {ext}")
            return None
        
        logger.info(f"Estrazione multimodale da: {path.name}")
        
        try:
            extractor = self.supported_formats[ext]
            doc = extractor(file_path)
            
            # Post-processing
            if doc:
                doc = self._consolidate_content(doc)
                doc = self._calculate_stats(doc)
            
            return doc
        except Exception as e:
            logger.error(f"Errore nell'estrazione di {path.name}: {e}")
            return None
    
    def extract_pdf_multimodal(self, file_path: str) -> ExtractedDocument:
        """Estrae testo, immagini e tabelle da PDF"""
        if not fitz:
            # Fallback a estrazione base
            return self._extract_pdf_basic(file_path)
        
        doc_obj = ExtractedDocument(
            source_path=file_path,
            file_type='pdf',
            title=Path(file_path).stem,
            content=""
        )
        
        # Apri con PyMuPDF
        pdf_document = fitz.open(file_path)
        
        # Estrai metadata
        metadata = pdf_document.metadata
        if metadata:
            doc_obj.metadata['pdf_metadata'] = metadata
            if metadata.get('title'):
                doc_obj.title = metadata['title']
        
        all_text_parts = []
        
        for page_num, page in enumerate(pdf_document, 1):
            logger.info(f"  Processing pagina {page_num}/{len(pdf_document)}")
            
            # Estrai testo
            page_text = page.get_text()
            
            # Estrai tabelle
            if self.extract_tables:
                tables = self._extract_tables_from_page(page, page_num)
                for table in tables:
                    doc_obj.tables.append(table)
                    # Inserisci placeholder nel testo
                    placeholder = f"\n[TABELLA_{table['table_id']}]\n"
                    page_text += placeholder
                    
                    # Aggiungi come elemento
                    elem = ExtractedElement(
                        element_type='table',
                        content=table['content'],
                        page_number=page_num,
                        metadata={'rows': table['rows'], 'cols': table['cols']}
                    )
                    doc_obj.elements.append(elem)
            
            # Estrai immagini
            if self.extract_images:
                images = self._extract_images_from_page(page, page_num, Path(file_path).stem)
                for img in images:
                    doc_obj.images.append(img)
                    # Inserisci placeholder nel testo
                    placeholder = f"\n[IMMAGINE_{img['image_id']}]\n"
                    page_text += placeholder
                    
                    # Aggiungi come elemento
                    elem = ExtractedElement(
                        element_type='image',
                        content=img['path'] if self.save_images else img['data'],
                        page_number=page_num,
                        metadata={'width': img['width'], 'height': img['height']}
                    )
                    doc_obj.elements.append(elem)
            
            all_text_parts.append(page_text)
            
            # Crea sezione per la pagina
            doc_obj.sections.append({
                'title': f'Pagina {page_num}',
                'content': page_text,
                'page_num': page_num
            })
        
        pdf_document.close()
        
        # Combina tutto il testo
        doc_obj.content = '\n\n'.join(all_text_parts)
        
        # Conta elementi
        doc_obj.image_count = len(doc_obj.images)
        doc_obj.table_count = len(doc_obj.tables)
        
        return doc_obj
    
    def _extract_tables_from_page(self, page, page_num: int) -> List[Dict]:
        """Estrae tabelle da una pagina PDF"""
        tables = []
        
        # Usa PyMuPDF per trovare tabelle (metodo semplificato)
        # In produzione, potresti usare librerie specializzate come camelot-py o tabula-py
        try:
            # Estrai blocchi di testo che sembrano tabelle
            blocks = page.get_text("blocks")
            
            for block_num, block in enumerate(blocks):
                # Identifica potenziali tabelle basandosi su pattern
                if self._looks_like_table(block):
                    table_data = self._parse_table_block(block)
                    if table_data:
                        table_id = f"p{page_num}_t{len(tables)+1}"
                        tables.append({
                            'table_id': table_id,
                            'page': page_num,
                            'content': table_data,
                            'rows': len(table_data),
                            'cols': len(table_data[0]) if table_data else 0,
                            'text_representation': self._table_to_text(table_data)
                        })
        except Exception as e:
            logger.warning(f"Errore estrazione tabelle pagina {page_num}: {e}")
        
        return tables
    
    def _extract_images_from_page(self, page, page_num: int, doc_name: str) -> List[Dict]:
        """Estrae immagini da una pagina PDF"""
        images = []
        
        try:
            image_list = page.get_images()
            
            for img_index, img in enumerate(image_list):
                try:
                    # Estrai immagine
                    xref = img[0]
                    pix = fitz.Pixmap(page.parent, xref)
                    
                    # Converti in PIL Image
                    if pix.n - pix.alpha < 4:  # GRAY o RGB
                        img_data = pix.tobytes("png")
                        img_pil = Image.open(BytesIO(img_data))
                    else:  # CMYK
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                        img_data = pix.tobytes("png")
                        img_pil = Image.open(BytesIO(img_data))
                    
                    # Genera ID e nome file
                    img_id = f"p{page_num}_img{img_index+1}"
                    img_filename = f"{doc_name}_{img_id}.png"
                    
                    img_info = {
                        'image_id': img_id,
                        'page': page_num,
                        'width': img_pil.width,
                        'height': img_pil.height,
                        'format': 'PNG'
                    }
                    
                    # Salva o conserva in memoria
                    if self.save_images:
                        img_path = self.images_dir / img_filename
                        img_pil.save(img_path)
                        img_info['path'] = str(img_path)
                        logger.info(f"    Salvata immagine: {img_filename}")
                    else:
                        # Conserva come base64
                        buffered = BytesIO()
                        img_pil.save(buffered, format="PNG")
                        img_info['data'] = base64.b64encode(buffered.getvalue()).decode()
                    
                    # OCR opzionale
                    if self.ocr_images and pytesseract:
                        try:
                            ocr_text = pytesseract.image_to_string(img_pil)
                            img_info['ocr_text'] = ocr_text
                        except Exception as e:
                            logger.warning(f"OCR fallito per {img_id}: {e}")
                    
                    images.append(img_info)
                    
                except Exception as e:
                    logger.warning(f"Errore estrazione immagine {img_index} pagina {page_num}: {e}")
                finally:
                    if pix:
                        pix = None
        
        except Exception as e:
            logger.warning(f"Errore estrazione immagini pagina {page_num}: {e}")
        
        return images
    
    def extract_docx_multimodal(self, file_path: str) -> ExtractedDocument:
        """Estrae contenuto multimodale da file Word"""
        if not DocxDocument:
            return self.extract_text(file_path)
        
        doc = DocxDocument(file_path)
        doc_obj = ExtractedDocument(
            source_path=file_path,
            file_type='docx',
            title=Path(file_path).stem,
            content=""
        )
        
        all_text_parts = []
        
        # Estrai paragrafi
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                all_text_parts.append(text)
                
                # Identifica titoli
                if para.style.name.startswith('Heading'):
                    doc_obj.sections.append({
                        'title': text,
                        'style': para.style.name
                    })
        
        # Estrai tabelle
        if self.extract_tables:
            for table_idx, table in enumerate(doc.tables):
                table_data = self._extract_docx_table(table)
                table_id = f"table_{table_idx+1}"
                
                doc_obj.tables.append({
                    'table_id': table_id,
                    'content': table_data,
                    'rows': len(table_data),
                    'cols': len(table_data[0]) if table_data else 0,
                    'text_representation': self._table_to_text(table_data)
                })
                
                # Aggiungi placeholder
                all_text_parts.append(f"\n[TABELLA_{table_id}]\n")
                
                # Aggiungi come elemento
                elem = ExtractedElement(
                    element_type='table',
                    content=table_data,
                    metadata={'rows': len(table_data), 'cols': len(table_data[0]) if table_data else 0}
                )
                doc_obj.elements.append(elem)
        
        # Estrai immagini incorporate
        if self.extract_images:
            # Word memorizza immagini nelle relazioni del documento
            try:
                for rel in doc.part.rels.values():
                    if "image" in rel.reltype:
                        img_data = rel.target_part.blob
                        img = Image.open(BytesIO(img_data))
                        
                        img_id = f"img_{len(doc_obj.images)+1}"
                        img_info = {
                            'image_id': img_id,
                            'width': img.width,
                            'height': img.height,
                            'format': img.format
                        }
                        
                        if self.save_images:
                            img_filename = f"{Path(file_path).stem}_{img_id}.png"
                            img_path = self.images_dir / img_filename
                            img.save(img_path)
                            img_info['path'] = str(img_path)
                        else:
                            buffered = BytesIO()
                            img.save(buffered, format="PNG")
                            img_info['data'] = base64.b64encode(buffered.getvalue()).decode()
                        
                        doc_obj.images.append(img_info)
                        all_text_parts.append(f"\n[IMMAGINE_{img_id}]\n")
            except Exception as e:
                logger.warning(f"Errore estrazione immagini DOCX: {e}")
        
        doc_obj.content = '\n\n'.join(all_text_parts)
        doc_obj.image_count = len(doc_obj.images)
        doc_obj.table_count = len(doc_obj.tables)
        
        return doc_obj
    
    def _extract_docx_table(self, table) -> List[List[str]]:
        """Estrae dati da una tabella Word"""
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data.append(row_data)
        return table_data
    
    def _looks_like_table(self, block) -> bool:
        """Euristica per identificare se un blocco sembra una tabella"""
        if len(block) < 5:  # block ha almeno 5 elementi
            return False
        
        text = str(block[4]) if len(block) > 4 else ""
        
        # Cerca pattern tipici delle tabelle
        if '\t' in text or '|' in text:
            return True
        
        # Conta spazi multipli che potrebbero indicare colonne
        if re.search(r'  {2,}', text):
            lines = text.split('\n')
            if len(lines) > 2:
                # Verifica allineamento
                spaces_per_line = [len(re.findall(r'  +', line)) for line in lines[:5]]
                if spaces_per_line and all(s > 0 for s in spaces_per_line):
                    return True
        
        return False
    
    def _parse_table_block(self, block) -> List[List[str]]:
        """Parsing semplice di un blocco che sembra una tabella"""
        if len(block) < 5:
            return []
        
        text = str(block[4])
        lines = text.split('\n')
        table_data = []
        
        for line in lines:
            if not line.strip():
                continue
            
            # Prova diversi delimitatori
            if '\t' in line:
                cells = line.split('\t')
            elif '|' in line:
                cells = [c.strip() for c in line.split('|') if c.strip()]
            else:
                # Usa spazi multipli come delimitatore
                cells = re.split(r'  +', line)
            
            if cells:
                table_data.append([c.strip() for c in cells])
        
        return table_data
    
    def _table_to_text(self, table_data: List[List[str]]) -> str:
        """Converte tabella in rappresentazione testuale"""
        if not table_data:
            return ""
        
        if pd and tabulate:
            try:
                df = pd.DataFrame(table_data[1:], columns=table_data[0] if table_data else [])
                return tabulate.tabulate(df, headers='keys', tablefmt='grid')
            except:
                pass
        
        # Fallback: rappresentazione semplice
        lines = []
        for row in table_data:
            lines.append(' | '.join(str(cell) for cell in row))
        return '\n'.join(lines)
    
    def _consolidate_content(self, doc: ExtractedDocument) -> ExtractedDocument:
        """Consolida il contenuto sostituendo placeholder con descrizioni"""
        content = doc.content
        
        # Sostituisci placeholder immagini con descrizioni
        for img in doc.images:
            placeholder = f"[IMMAGINE_{img['image_id']}]"
            description = f"[Immagine: {img['image_id']} - Dimensioni: {img['width']}x{img['height']}px"
            if 'ocr_text' in img and img['ocr_text']:
                description += f" - Testo OCR: {img['ocr_text'][:100]}..."
            description += "]"
            content = content.replace(placeholder, description)
        
        # Sostituisci placeholder tabelle con rappresentazione testuale
        for table in doc.tables:
            placeholder = f"[TABELLA_{table['table_id']}]"
            description = f"[Tabella {table['table_id']}: {table['rows']} righe x {table['cols']} colonne]\n"
            description += table.get('text_representation', '')
            content = content.replace(placeholder, description)
        
        doc.content = content
        return doc
    
    def _calculate_stats(self, doc: ExtractedDocument) -> ExtractedDocument:
        """Calcola statistiche del documento"""
        doc.char_count = len(doc.content)
        doc.word_count = len(doc.content.split())
        doc.line_count = doc.content.count('\n') + 1
        return doc
    
    # Metodi di fallback per altri formati
    def extract_text(self, file_path: str) -> ExtractedDocument:
        """Estrae contenuto da file di testo"""
        encoding = 'utf-8'
        if chardet:
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                result = chardet.detect(raw_data)
                encoding = result.get('encoding', 'utf-8')
        
        with open(file_path, 'r', encoding=encoding) as file:
            content = file.read()
        
        return ExtractedDocument(
            source_path=file_path,
            file_type='text',
            title=Path(file_path).stem,
            content=content,
            metadata={'encoding': encoding}
        )
    
    def extract_markdown(self, file_path: str) -> ExtractedDocument:
        """Estrae contenuto da Markdown"""
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        # Estrai titolo e sezioni
        title_match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
        title = title_match.group(1) if title_match else Path(file_path).stem
        
        sections = []
        section_pattern = r'^(#{1,6})\s+(.+)$'
        
        for match in re.finditer(section_pattern, content, re.MULTILINE):
            level = len(match.group(1))
            section_title = match.group(2)
            sections.append({
                'title': section_title,
                'level': level,
                'position': match.start()
            })
        
        # Identifica immagini Markdown
        images = []
        img_pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
        for match in re.finditer(img_pattern, content):
            alt_text = match.group(1)
            img_path = match.group(2)
            img_id = f"md_img_{len(images)+1}"
            images.append({
                'image_id': img_id,
                'alt_text': alt_text,
                'path': img_path
            })
        
        return ExtractedDocument(
            source_path=file_path,
            file_type='markdown',
            title=title,
            content=content,
            sections=sections,
            images=images,
            image_count=len(images)
        )
    
    def extract_html(self, file_path: str) -> ExtractedDocument:
        """Estrae contenuto da HTML"""
        with open(file_path, 'r', encoding='utf-8') as file:
            html_content = file.read()
        
        doc = ExtractedDocument(
            source_path=file_path,
            file_type='html',
            title=Path(file_path).stem,
            content=""
        )
        
        if BeautifulSoup:
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Rimuovi script e style
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Estrai titolo
            title = soup.find('title')
            doc.title = title.get_text() if title else Path(file_path).stem
            
            # Estrai immagini
            if self.extract_images:
                for img in soup.find_all('img'):
                    img_id = f"html_img_{len(doc.images)+1}"
                    doc.images.append({
                        'image_id': img_id,
                        'src': img.get('src', ''),
                        'alt': img.get('alt', ''),
                        'width': img.get('width', ''),
                        'height': img.get('height', '')
                    })
            
            # Estrai tabelle
            if self.extract_tables:
                for table_idx, table in enumerate(soup.find_all('table')):
                    table_data = self._parse_html_table(table)
                    if table_data:
                        table_id = f"html_table_{table_idx+1}"
                        doc.tables.append({
                            'table_id': table_id,
                            'content': table_data,
                            'rows': len(table_data),
                            'cols': len(table_data[0]) if table_data else 0
                        })
            
            # Estrai testo
            doc.content = soup.get_text()
            
            # Estrai sezioni da heading
            for heading in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                doc.sections.append({
                    'title': heading.get_text().strip(),
                    'level': int(heading.name[1])
                })
        else:
            # Fallback senza BeautifulSoup
            doc.content = re.sub(r'<[^>]+>', '', html_content)
        
        doc.image_count = len(doc.images)
        doc.table_count = len(doc.tables)
        
        return doc
    
    def _parse_html_table(self, table) -> List[List[str]]:
        """Estrae dati da una tabella HTML"""
        table_data = []
        
        # Header
        headers = []
        for th in table.find_all('th'):
            headers.append(th.get_text().strip())
        if headers:
            table_data.append(headers)
        
        # Righe
        for tr in table.find_all('tr'):
            row = []
            for td in tr.find_all('td'):
                row.append(td.get_text().strip())
            if row:
                table_data.append(row)
        
        return table_data
    
    def extract_json(self, file_path: str) -> ExtractedDocument:
        """Estrae contenuto da JSON"""
        with open(file_path, 'r', encoding='utf-8') as file:
            data = json.load(file)
        
        content = json.dumps(data, indent=2, ensure_ascii=False)
        
        return ExtractedDocument(
            source_path=file_path,
            file_type='json',
            title=Path(file_path).stem,
            content=content,
            metadata={'json_structure': type(data).__name__}
        )
    
    def _extract_pdf_basic(self, file_path: str) -> ExtractedDocument:
        """Fallback: estrazione base da PDF senza PyMuPDF"""
        if not PdfReader:
            raise ImportError("pypdf richiesto per leggere PDF")
        
        with open(file_path, 'rb') as file:
            reader = PdfReader(file)
            
            pages_text = []
            for i, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                if page_text.strip():
                    pages_text.append(page_text)
            
            content = '\n\n'.join(pages_text)
            
            return ExtractedDocument(
                source_path=file_path,
                file_type='pdf',
                title=Path(file_path).stem,
                content=content,
                metadata={'page_count': len(reader.pages)}
            )


# Mantieni TextCleaner dalla versione precedente
class TextCleaner:
    """Pulisce e normalizza il testo estratto"""
    
    def __init__(self, aggressive: bool = False, preserve_structure: bool = True):
        self.aggressive = aggressive
        self.preserve_structure = preserve_structure
    
    def clean(self, text: str, doc_type: str = 'generic') -> str:
        """Pipeline di pulizia del testo preservando riferimenti a elementi multimodali"""
        original_len = len(text)
        
        # Proteggi i placeholder di immagini e tabelle
        protected_patterns = []
        
        # Trova e proteggi placeholder
        img_pattern = r'\[IMMAGINE_[^\]]+\]'
        table_pattern = r'\[TABELLA_[^\]]+\]'
        
        for pattern in [img_pattern, table_pattern]:
            matches = re.findall(pattern, text)
            protected_patterns.extend(matches)
        
        # Sostituisci temporaneamente con marcatori unici
        for i, pattern in enumerate(protected_patterns):
            text = text.replace(pattern, f"<<<PROTECTED_{i}>>>")
        
        # Pipeline di pulizia standard
        text = self._normalize_unicode(text)
        text = self._fix_encoding_errors(text)
        text = self._remove_control_characters(text)
        text = self._normalize_whitespace(text)
        text = self._fix_hyphenation(text)
        text = self._normalize_punctuation(text)
        
        if doc_type == 'pdf':
            text = self._clean_pdf_artifacts(text)
        
        if self.aggressive:
            text = self._aggressive_clean(text)
        
        # Ripristina placeholder protetti
        for i, pattern in enumerate(protected_patterns):
            text = text.replace(f"<<<PROTECTED_{i}>>>", pattern)
        
        cleaned_len = len(text)
        reduction = ((original_len - cleaned_len) / original_len * 100) if original_len > 0 else 0
        logger.info(f"Testo pulito: {original_len} -> {cleaned_len} caratteri ({reduction:.1f}% riduzione)")
        
        return text
    
    # [Mantieni tutti i metodi di pulizia dalla versione precedente]
    def _normalize_unicode(self, text: str) -> str:
        return unicodedata.normalize('NFKC', text)
    
    def _fix_encoding_errors(self, text: str) -> str:
        replacements = {
            'â€™': "'", 'â€œ': '"', 'â€': '"', 'â€"': '—', 'â€"': '–',
            'Ã¨': 'è', 'Ã©': 'é', 'Ã': 'à', 'Ã¹': 'ù', 'Ã²': 'ò', 'â‚¬': '€',
        }
        for wrong, correct in replacements.items():
            text = text.replace(wrong, correct)
        return text
    
    def _remove_control_characters(self, text: str) -> str:
        return ''.join(char for char in text if char in '\t\n' or not unicodedata.category(char).startswith('C'))
    
    def _normalize_whitespace(self, text: str) -> str:
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n{3,}', '\n\n', text)
        lines = text.split('\n')
        lines = [line.strip() for line in lines]
        return '\n'.join(lines)
    
    def _fix_hyphenation(self, text: str) -> str:
        text = re.sub(r'([a-z])-\n([a-z])', r'\1\2', text)
        text = re.sub(r'([a-z])-\n([A-Z])', r'\1-\n\2', text)
        return text
    
    def _normalize_punctuation(self, text: str) -> str:
        text = re.sub(r'([.!?,:;])([A-Za-z])', r'\1 \2', text)
        text = re.sub(r'\s+([.!?,:;])', r'\1', text)
        text = re.sub(r'[\u201C\u201D]', '"', text)  # Replace smart quotes with straight quotes
        text = re.sub(r'[\u2018\u2019]', "'", text)  # Replace smart single quotes with straight single quote
        text = re.sub(r'\.{3,}', '...', text)
        return text
    
    def _clean_pdf_artifacts(self, text: str) -> str:
        text = re.sub(r'^\d{1,4}\s*$', '', text, flags=re.MULTILINE)
        text = re.sub(r'^Page \d+ of \d+\s*$', '', text, flags=re.MULTILINE)
        text = re.sub(r'^[_\-=*]{3,}\s*$', '', text, flags=re.MULTILINE)
        return text
    
    def _aggressive_clean(self, text: str) -> str:
        text = re.sub(r'https?://\S+', '[URL]', text)
        text = re.sub(r'\S+@\S+\.\S+', '[EMAIL]', text)
        text = re.sub(r'\b[A-Z0-9]{10,}\b', '[CODE]', text)
        lines = text.split('\n')
        lines = [line for line in lines if len(line) > 10 or line == '']
        return '\n'.join(lines)


class MultimodalDocumentPreprocessor:
    """Preprocessore principale con supporto multimodale"""
    
    def __init__(self, 
                 output_dir: str = "./processed",
                 images_dir: str = "./extracted_images",
                 extract_images: bool = True,
                 extract_tables: bool = True,
                 ocr_images: bool = False,
                 aggressive_clean: bool = False):
        
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        
        self.extractor = MultimodalDocumentExtractor(
            extract_images=extract_images,
            extract_tables=extract_tables,
            ocr_images=ocr_images,
            save_images=True,
            images_dir=images_dir
        )
        
        self.cleaner = TextCleaner(
            aggressive=aggressive_clean,
            preserve_structure=True
        )
        
        self.processed_documents = []
        self.all_images = []
        self.all_tables = []
    
    def process_directory(self, 
                         input_dir: str,
                         recursive: bool = True,
                         output_file: str = "consolidated_multimodal.txt",
                         save_manifest: bool = True) -> str:
        """
        Processa tutti i documenti in una directory con supporto multimodale
        """
        input_path = Path(input_dir)
        
        if not input_path.exists():
            raise ValueError(f"Directory non trovata: {input_dir}")
        
        logger.info(f"Processing multimodale directory: {input_path}")
        
        # Trova tutti i file
        if recursive:
            files = list(input_path.rglob("*"))
        else:
            files = list(input_path.glob("*"))
        
        # Filtra solo file supportati
        supported_extensions = self.extractor.supported_formats.keys()
        files = [f for f in files if f.is_file() and f.suffix.lower() in supported_extensions]
        
        logger.info(f"Trovati {len(files)} file da processare")
        
        # Processa ogni file
        for file_path in files:
            doc = self.process_file(str(file_path))
            if doc:
                self.all_images.extend(doc.images)
                self.all_tables.extend(doc.tables)
        
        # Consolida tutti i documenti
        output_path = self.output_dir / output_file
        self.consolidate_documents(str(output_path))
        
        # Salva manifest degli elementi multimodali
        if save_manifest:
            self._save_multimodal_manifest()
        
        return str(output_path)
    
    def process_file(self, file_path: str) -> Optional[ExtractedDocument]:
        """Processa un singolo file con estrazione multimodale"""
        logger.info(f"Processing multimodale: {Path(file_path).name}")
        
        # Estrai contenuto multimodale
        doc = self.extractor.extract(file_path)
        
        if not doc:
            logger.warning(f"Impossibile estrarre contenuto da: {file_path}")
            return None
        
        # Pulisci contenuto testuale preservando riferimenti
        doc.content = self.cleaner.clean(doc.content, doc.file_type)
        
        # Aggiungi alla lista dei documenti processati
        self.processed_documents.append(doc)
        
        logger.info(f"✓ Processato: {doc.title}")
        logger.info(f"  - Parole: {doc.word_count:,}")
        logger.info(f"  - Immagini: {doc.image_count}")
        logger.info(f"  - Tabelle: {doc.table_count}")
        
        return doc
    
    def consolidate_documents(self, output_file: str):
        """
        Consolida tutti i documenti processati con riferimenti multimodali
        """
        if not self.processed_documents:
            logger.warning("Nessun documento da consolidare")
            return
        
        logger.info(f"Consolidamento di {len(self.processed_documents)} documenti...")
        
        # Ordina documenti
        docs = sorted(self.processed_documents, 
                     key=lambda d: (d.file_type != 'pdf', -d.word_count))
        
        # Crea contenuto consolidato
        consolidated_lines = []
        
        # Header
        consolidated_lines.append("=" * 80)
        consolidated_lines.append("DOCUMENTI CONSOLIDATI CON ELEMENTI MULTIMODALI")
        consolidated_lines.append(f"Data: {datetime.now().isoformat()}")
        consolidated_lines.append(f"Totale documenti: {len(docs)}")
        consolidated_lines.append(f"Totale immagini: {sum(d.image_count for d in docs)}")
        consolidated_lines.append(f"Totale tabelle: {sum(d.table_count for d in docs)}")
        consolidated_lines.append("=" * 80)
        consolidated_lines.append("")
        
        # Indice
        consolidated_lines.append("INDICE DEI DOCUMENTI:")
        consolidated_lines.append("-" * 40)
        for i, doc in enumerate(docs, 1):
            consolidated_lines.append(f"{i}. {doc.title}")
            stats = f"   Tipo: {doc.file_type} | Parole: {doc.word_count:,}"
            if doc.image_count > 0:
                stats += f" | Immagini: {doc.image_count}"
            if doc.table_count > 0:
                stats += f" | Tabelle: {doc.table_count}"
            consolidated_lines.append(stats)
        consolidated_lines.append("")
        
        # Contenuto dei documenti
        for i, doc in enumerate(docs, 1):
            consolidated_lines.append("")
            consolidated_lines.append("=" * 80)
            consolidated_lines.append(f"DOCUMENTO {i}: {doc.title}")
            consolidated_lines.append(f"Fonte: {Path(doc.source_path).name}")
            consolidated_lines.append(f"Tipo: {doc.file_type}")
            consolidated_lines.append(f"Statistiche: Parole: {doc.word_count:,} | Caratteri: {doc.char_count:,}")
            
            if doc.image_count > 0 or doc.table_count > 0:
                consolidated_lines.append(f"Elementi multimodali: Immagini: {doc.image_count} | Tabelle: {doc.table_count}")
            
            consolidated_lines.append("=" * 80)
            consolidated_lines.append("")
            
            # Contenuto principale con riferimenti multimodali
            consolidated_lines.append("CONTENUTO:")
            consolidated_lines.append("-" * 40)
            consolidated_lines.append(doc.content)
            consolidated_lines.append("")
            
            # Lista elementi multimodali
            if doc.images:
                consolidated_lines.append("\nIMMAGINI NEL DOCUMENTO:")
                for img in doc.images[:10]:  # Prime 10 immagini
                    img_desc = f"  - {img['image_id']}: {img.get('width', 'N/A')}x{img.get('height', 'N/A')}px"
                    if 'path' in img:
                        img_desc += f" - Salvata in: {Path(img['path']).name}"
                    consolidated_lines.append(img_desc)
            
            if doc.tables:
                consolidated_lines.append("\nTABELLE NEL DOCUMENTO:")
                for table in doc.tables[:10]:  # Prime 10 tabelle
                    consolidated_lines.append(f"  - {table['table_id']}: {table['rows']}x{table['cols']}")
            
            consolidated_lines.append("")
            consolidated_lines.append(f"[Fine documento {i}]")
            consolidated_lines.append("")
        
        # Footer
        consolidated_lines.append("")
        consolidated_lines.append("=" * 80)
        consolidated_lines.append("RIEPILOGO FINALE")
        consolidated_lines.append(f"Documenti: {len(docs)}")
        consolidated_lines.append(f"Parole totali: {sum(d.word_count for d in docs):,}")
        consolidated_lines.append(f"Caratteri totali: {sum(d.char_count for d in docs):,}")
        consolidated_lines.append(f"Immagini totali: {len(self.all_images)}")
        consolidated_lines.append(f"Tabelle totali: {len(self.all_tables)}")
        consolidated_lines.append("=" * 80)
        
        # Scrivi file
        full_content = '\n'.join(consolidated_lines)
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(full_content)
        
        logger.info(f"✓ File consolidato salvato: {output_file}")
        logger.info(f"  Dimensione: {len(full_content):,} caratteri")
        
        # Salva metadata in JSON
        self._save_metadata(output_file)
    
    def _save_metadata(self, output_file: str):
        """Salva metadata completi in JSON"""
        metadata_file = Path(output_file).with_suffix('.json')
        
        metadata = {
            'creation_date': datetime.now().isoformat(),
            'total_documents': len(self.processed_documents),
            'total_words': sum(d.word_count for d in self.processed_documents),
            'total_chars': sum(d.char_count for d in self.processed_documents),
            'total_images': len(self.all_images),
            'total_tables': len(self.all_tables),
            'documents': [
                {
                    'title': d.title,
                    'source': d.source_path,
                    'type': d.file_type,
                    'words': d.word_count,
                    'chars': d.char_count,
                    'images': d.image_count,
                    'tables': d.table_count,
                    'sections_count': len(d.sections)
                }
                for d in self.processed_documents
            ]
        }
        
        with open(metadata_file, 'w', encoding='utf-8') as f:
            json.dump(metadata, f, indent=2, ensure_ascii=False)
        
        logger.info(f"✓ Metadata salvati: {metadata_file}")
    
    def _save_multimodal_manifest(self):
        """Salva manifest degli elementi multimodali"""
        manifest_file = self.output_dir / "multimodal_manifest.json"
        
        manifest = {
            'creation_date': datetime.now().isoformat(),
            'images': [
                {
                    'id': img['image_id'],
                    'path': img.get('path', ''),
                    'width': img.get('width', 0),
                    'height': img.get('height', 0),
                    'source_doc': img.get('source_doc', '')
                }
                for img in self.all_images
            ],
            'tables': [
                {
                    'id': table['table_id'],
                    'rows': table.get('rows', 0),
                    'cols': table.get('cols', 0),
                    'source_doc': table.get('source_doc', '')
                }
                for table in self.all_tables
            ]
        }
        
        with open(manifest_file, 'w', encoding='utf-8') as f:
            json.dump(manifest, f, indent=2, ensure_ascii=False)
        
        logger.info(f"✓ Manifest multimodale salvato: {manifest_file}")


def main():
    """Funzione principale"""
    import argparse
    
    parser = argparse.ArgumentParser(description='Preprocessa documenti con supporto multimodale')
    parser.add_argument('input_dir', help='Directory contenente i documenti')
    parser.add_argument('-o', '--output', default='./processed', help='Directory output')
    parser.add_argument('-i', '--images', default='./extracted_images', help='Directory immagini')
    parser.add_argument('-f', '--file', default='consolidated_multimodal.txt', help='Nome file output')
    parser.add_argument('-r', '--recursive', action='store_true', help='Scansione ricorsiva')
    parser.add_argument('-a', '--aggressive', action='store_true', help='Pulizia aggressiva')
    parser.add_argument('--no-images', action='store_true', help='Non estrarre immagini')
    parser.add_argument('--no-tables', action='store_true', help='Non estrarre tabelle')
    parser.add_argument('--ocr', action='store_true', help='Applica OCR alle immagini')
    
    args = parser.parse_args()
    
    # Crea preprocessore multimodale
    preprocessor = MultimodalDocumentPreprocessor(
        output_dir=args.output,
        images_dir=args.images,
        extract_images=not args.no_images,
        extract_tables=not args.no_tables,
        ocr_images=args.ocr,
        aggressive_clean=args.aggressive
    )
    
    # Processa directory
    try:
        output_file = preprocessor.process_directory(
            args.input_dir,
            recursive=args.recursive,
            output_file=args.file
        )
        
        print(f"\n✅ Elaborazione multimodale completata!")
        print(f"📄 File consolidato: {output_file}")
        print(f"📊 Documenti: {len(preprocessor.processed_documents)}")
        print(f"🖼️  Immagini estratte: {len(preprocessor.all_images)}")
        print(f"📋 Tabelle estratte: {len(preprocessor.all_tables)}")
        
    except Exception as e:
        logger.error(f"Errore durante l'elaborazione: {e}")
        return 1
    
    return 0


if __name__ == "__main__":
    # Esempio di utilizzo diretto
    INPUT_DIR = "../documents"
    OUTPUT_DIR = "./processed"
    IMAGES_DIR = "./extracted_images"
    
    # Crea preprocessore multimodale
    preprocessor = MultimodalDocumentPreprocessor(
        output_dir=OUTPUT_DIR,
        images_dir=IMAGES_DIR,
        extract_images=True,
        extract_tables=True,
        ocr_images=False,  # Attiva se hai pytesseract configurato
        aggressive_clean=False
    )
    
    # Processa tutti i documenti
    output_file = preprocessor.process_directory(
        INPUT_DIR,
        recursive=True,
        output_file="documents_multimodal.txt",
        save_manifest=True
    )
    
    print(f"\n✅ Processo completato!")
    print(f"📄 File generato: {output_file}")
    print(f"📊 Documenti: {len(preprocessor.processed_documents)}")
    print(f"🖼️  Immagini: {len(preprocessor.all_images)}")
    print(f"📋 Tabelle: {len(preprocessor.all_tables)}")